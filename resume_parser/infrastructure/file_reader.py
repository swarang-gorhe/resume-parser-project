from __future__ import annotations

import subprocess
import tempfile
from abc import ABC, abstractmethod
from pathlib import Path

import docx
import pdfplumber

try:
    import magic
except ImportError:  # pragma: no cover
    magic = None

from resume_parser.core.exceptions import ConversionError, UnsupportedFileTypeError
from resume_parser.core.logging import get_logger


class BaseFileReader(ABC):
    def __init__(self, file_path: Path) -> None:
        self.file_path = file_path
        self.logger = get_logger(self.__class__.__name__)

    @abstractmethod
    def read_text(self) -> str:
        pass


class PDFReader(BaseFileReader):
    def read_text(self) -> str:
        self.logger.debug("Reading PDF file.", extra={"event": "file_reader.pdf_read"})
        with pdfplumber.open(self.file_path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages).strip()


class DOCXReader(BaseFileReader):
    def read_text(self) -> str:
        self.logger.debug("Reading DOCX file.", extra={"event": "file_reader.docx_read"})
        document = docx.Document(self.file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
        return "\n".join(paragraphs + table_cells).strip()


class DOCReader(BaseFileReader):
    def read_text(self) -> str:
        self.logger.debug("Converting DOC file to DOCX.", extra={"event": "file_reader.doc_conversion"})
        with tempfile.TemporaryDirectory() as tmpdir:
            output_dir = Path(tmpdir)
            try:
                subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        str(self.file_path),
                        "--outdir",
                        str(output_dir),
                    ],
                    check=True,
                    capture_output=True,
                )
            except subprocess.CalledProcessError as exc:
                raise ConversionError("Failed to convert DOC to DOCX") from exc
            converted = output_dir / f"{self.file_path.stem}.docx"
            if not converted.exists():
                raise ConversionError("DOC conversion did not produce a DOCX file.")
            return DOCXReader(converted).read_text()


class FileReaderFactory:
    @staticmethod
    def get_reader(file_path: Path) -> BaseFileReader:
        mime = None
        if magic is not None:
            try:
                mime = magic.from_file(str(file_path), mime=True)
            except Exception:
                mime = None
        extension = file_path.suffix.lower()
        if mime == "application/pdf" or extension == ".pdf":
            return PDFReader(file_path)
        if mime in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"} or extension in {".docx", ".doc"}:
            if extension == ".docx":
                return DOCXReader(file_path)
            return DOCReader(file_path)
        raise UnsupportedFileTypeError(f"Unsupported MIME type: {mime or 'unknown'}")
